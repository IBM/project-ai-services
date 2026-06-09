"""
Utilities for processing DOCX files, providing page numbers, font sizes, and TOC.
This module provides DOCX-specific functionality parallel to PDF utilities.
"""
import json
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from docx import Document

from common.misc_utils import get_logger

logger = get_logger("docx_utils")

# Constants for page estimation
CHARS_PER_PAGE = 3000  # Approximate characters per page
WORDS_PER_PAGE = 450   # Approximate words per page


def estimate_docx_page_count(docx_path: str) -> int:
    """
    Estimate page count for DOCX based on content and formatting.
    Assumes standard page: ~450 words or ~3000 characters per page.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Estimated page count (minimum 1)
    """
    try:
        doc = Document(docx_path)
        
        total_chars = 0
        total_words = 0
        
        # Count paragraph content
        for para in doc.paragraphs:
            text = para.text
            total_chars += len(text)
            total_words += len(text.split())
        
        # Count table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    total_chars += len(text)
                    total_words += len(text.split())
        
        # Estimate pages (conservative: 450 words per page)
        estimated_pages = max(1, (total_words // WORDS_PER_PAGE) + 1)
        
        logger.debug(f"DOCX {docx_path}: {total_words} words, estimated {estimated_pages} pages")
        return estimated_pages
        
    except Exception as e:
        logger.error(f"Error estimating page count for {docx_path}: {e}")
        return 1  # Return minimum 1 page on error


def get_docx_toc(docx_path: str) -> Dict[str, int]:
    """
    Extract table of contents from DOCX file.
    
    Strategy (in order of preference):
    1. Try combined extraction (Table Paragraph + List Paragraph) - handles most TOC formats
    2. Fall back to formal TOC styles ('TOC 1', 'TOC 2', etc.)
    3. Fall back to Heading styles ('Heading 1', 'Heading 2', etc.)
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping heading text to level (similar to PDF TOC format)
    """
    try:
        # Primary method: Combined extraction (Table Paragraph + List Paragraph)
        toc = extract_toc_combined(docx_path)
        
        if toc:
            logger.info(f"DOCX {docx_path}: extracted {len(toc)} TOC entries from combined extraction")
            return toc
        
        # Fallback 1: Try formal TOC styles
        logger.info(f"DOCX {docx_path}: No combined TOC found, trying TOC styles")
        toc = extract_toc_from_toc_styles(docx_path)
        
        if toc:
            logger.info(f"DOCX {docx_path}: extracted {len(toc)} TOC entries from TOC styles")
            return toc
        
        # Fallback 2: Try Heading styles
        logger.info(f"DOCX {docx_path}: No TOC styles found, trying Heading styles")
        toc = extract_toc_from_headings(docx_path)
        
        if toc:
            logger.info(f"DOCX {docx_path}: extracted {len(toc)} TOC entries from Heading styles")
        else:
            logger.warning(f"DOCX {docx_path}: No TOC found using any method")

        logger.debug(f" extracted toc is {toc}")
        
        return toc
        
    except Exception as e:
        logger.error(f"Error extracting TOC from {docx_path}: {e}")
        return {}


def get_docx_header_font_sizes(docx_path: str) -> Dict[int, float]:
    """
    Extract font sizes for headers in DOCX file.
    Returns: Dict mapping heading levels to average font sizes
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping heading level to average font size in points
    """
    try:
        doc = Document(docx_path)
        header_font_sizes = {}
        
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith('Heading'):
                # Extract heading level
                match = re.match(r'Heading (\d+)', para.style.name)
                if match:
                    level = int(match.group(1))
                    
                    # Get font size from runs
                    if para.runs:
                        font_sizes = []
                        for run in para.runs:
                            if run.font.size:
                                # Convert to points
                                font_sizes.append(run.font.size.pt)
                        
                        if font_sizes:
                            avg_size = sum(font_sizes) / len(font_sizes)
                            if level not in header_font_sizes:
                                header_font_sizes[level] = []
                            header_font_sizes[level].append(avg_size)
        
        # Average font sizes per level
        result = {level: sum(sizes)/len(sizes) for level, sizes in header_font_sizes.items() if sizes}
        logger.debug(f"DOCX {docx_path}: extracted font sizes for {len(result)} heading levels")
        return result
        
    except Exception as e:
        logger.error(f"Error extracting header font sizes from {docx_path}: {e}")
        return {}


def assign_page_numbers_to_docx_content(docx_path: str) -> Tuple[List[Tuple[str, int, Dict]], int]:
    """
    Assign estimated page numbers to DOCX paragraphs based on content accumulation.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Tuple of (content_with_pages, total_pages) where content_with_pages is a list of
        (paragraph_text, estimated_page_number, style_info) tuples
    """
    try:
        doc = Document(docx_path)
        
        cumulative_chars = 0
        cumulative_words = 0
        current_page = 1
        
        content_with_pages = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            para_chars = len(text)
            para_words = len(text.split())
            
            # Assign current page
            page_number = current_page
            
            # Get style information
            style_info = {
                'style_name': para.style.name if para.style else None,
                'font_size': None,
                'is_heading': False,
                'heading_level': None
            }
            
            # Extract font size if available
            if para.runs:
                font_sizes = []
                for run in para.runs:
                    if run.font.size:
                        font_sizes.append(run.font.size.pt)
                if font_sizes:
                    style_info['font_size'] = max(font_sizes)  # Use largest font in paragraph
            
            # Check if it's a heading
            if para.style and para.style.name.startswith('Heading'):
                style_info['is_heading'] = True
                # Extract heading level (e.g., "Heading 1" -> 1)
                match = re.match(r'Heading (\d+)', para.style.name)
                if match:
                    style_info['heading_level'] = int(match.group(1))
            
            content_with_pages.append((text, page_number, style_info))
            
            # Update cumulative counts
            cumulative_chars += para_chars
            cumulative_words += para_words
            
            # Check if we should move to next page
            if cumulative_chars >= CHARS_PER_PAGE or cumulative_words >= WORDS_PER_PAGE:
                current_page += 1
                cumulative_chars = 0
                cumulative_words = 0
        
        logger.debug(f"DOCX {docx_path}: assigned page numbers to {len(content_with_pages)} paragraphs across {current_page} pages")
        return content_with_pages, current_page
        
    except Exception as e:
        logger.error(f"Error assigning page numbers for {docx_path}: {e}")
        return [], 1


def create_docx_page_mapping(docx_path: str) -> Dict[str, Tuple[int, Dict]]:
    """
    Create a mapping from text content to page numbers and style info for DOCX files.
    This is used to enrich Docling-converted content with page information.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping text content to (page_number, style_info) tuples
    """
    try:
        content_with_pages, _ = assign_page_numbers_to_docx_content(docx_path)
        
        # Create mapping from text to (page, style)
        page_map = {}
        for text, page, style in content_with_pages:
            # Use text as key (may have duplicates, last one wins)
            page_map[text] = (page, style)
        
        logger.debug(f"DOCX {docx_path}: created page mapping with {len(page_map)} entries")
        return page_map
        
    except Exception as e:
        logger.error(f"Error creating page mapping for {docx_path}: {e}")
        return {}



# ============================================================================
# NEW: TOC Style-Based Extraction Functions
# These functions extract TOC from Word's TOC field styles ('TOC 1', 'TOC 2', etc.)
# ============================================================================

def extract_toc_level_from_style(style_name: str) -> int:
    """
    Extract TOC level from Word style name.
    
    Examples:
        'TOC 1' -> 1
        'TOC 2' -> 2
        'toc 3' -> 3
        'TOC Heading' -> 1 (default)
    
    Args:
        style_name: Word paragraph style name
        
    Returns:
        TOC level as integer (defaults to 1 if no number found)
    """
    # Try to extract number from style name
    match = re.search(r'toc\s*(\d+)', style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    
    # If no number found (e.g., 'TOC Heading'), default to level 1
    return 1


def extract_toc_from_toc_styles(docx_path: str) -> Dict[str, int]:
    """
    Extract TOC from DOCX file by looking for TOC styles.
    This works ONLY if the document has a formal TOC field with TOC styles.
    
    Word TOC styles are typically named: 'TOC 1', 'TOC 2', 'TOC 3', 'TOC Heading', etc.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping TOC text to level (similar to PDF TOC format)
    """
    try:
        doc = Document(docx_path)
        toc = {}
        
        for paragraph in doc.paragraphs:
            if paragraph.style and paragraph.style.name:
                style_name = paragraph.style.name
                
                # Check for TOC styles: 'TOC 1', 'TOC 2', 'TOC Heading', etc.
                if 'toc' in style_name.lower():
                    text = paragraph.text.strip()
                    if text:
                        # Extract level from style name
                        level = extract_toc_level_from_style(style_name)
                        toc[text] = level
        
        logger.info(f"Extracted {len(toc)} TOC entries from TOC styles in {docx_path}")
        if len(toc) == 0:
            logger.warning(
                f"No TOC styles found in {docx_path}. "
                "Document may not have a formal TOC field. "
                "Consider using extract_toc_from_headings() instead."
            )
        
        return toc
        
    except Exception as e:
        logger.error(f"Error extracting TOC from TOC styles in {docx_path}: {e}")
        return {}


def extract_toc_from_headings(docx_path: str) -> Dict[str, int]:
    """
    Extract TOC from DOCX Heading styles (Heading 1, Heading 2, etc.).
    This is a fallback when document doesn't have formal TOC field.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping heading text to level
    """
    try:
        doc = Document(docx_path)
        toc = {}
        
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith('Heading'):
                text = para.text.strip()
                if text:
                    # Extract level: 'Heading 1' -> 1, 'Heading 2' -> 2
                    match = re.match(r'Heading (\d+)', para.style.name)
                    if match:
                        level = int(match.group(1))
                        toc[text] = level
        
        logger.info(f"Extracted {len(toc)} TOC entries from Heading styles in {docx_path}")
        return toc
        
    except Exception as e:
        logger.error(f"Error extracting TOC from Heading styles in {docx_path}: {e}")
        return {}


def extract_toc_combined(docx_path: str) -> Dict[str, int]:
    """
    Extract TOC from BOTH 'Table Paragraph' and 'List Paragraph' styles.
    This captures TOC entries whether they're in tables or as list items.
    
    This is the recommended approach as it handles various TOC formats:
    - TOC in tables (Table Paragraph style)
    - TOC as lists (List Paragraph style)
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping TOC text to level
    """
    try:
        doc = Document(docx_path)
        toc = {}
        
        logger.debug(f"Extracting TOC from Table Paragraphs AND List Paragraphs in {docx_path}")
        
        # Method 1: Extract from tables (Table Paragraph style)
        table_count = 0
        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.style and para.style.name == 'Table Paragraph':
                            text = para.text.strip()
                            if text:
                                # Skip header row (usually "Contents")
                                if text.lower() in ['contents', 'table of contents']:
                                    continue
                                
                                # Remove page numbers (e.g., "Chapter 1    45" -> "Chapter 1")
                                text_clean = re.sub(r'\s+\d+$', '', text).strip()
                                # Remove trailing dots/leader characters (e.g., "Chapter 1. . . . ." -> "Chapter 1")
                                # Pattern matches dots with spaces between them: ". . . . ." or "....."
                                text_clean = re.sub(r'[\.\s]+$', '', text_clean).strip()
                                
                                if text_clean and text_clean not in toc:
                                    # Skip standalone numbers (likely page numbers or section markers)
                                    if re.match(r'^\d+$', text_clean):
                                        continue
                                    level = _infer_toc_level_from_text(text_clean)
                                    toc[text_clean] = level
                                    table_count += 1
        
        logger.debug(f"Extracted {table_count} entries from Table Paragraph style")
        
        # Method 2: Extract from paragraphs (List Paragraph style)
        list_count = 0
        for para in doc.paragraphs:
            if para.style and para.style.name == 'List Paragraph':
                text = para.text.strip()
                if text:
                    # Skip header
                    if text.lower() in ['contents', 'table of contents']:
                        continue
                    
                    # Remove page numbers
                    text_clean = re.sub(r'\s+\d+$', '', text).strip()
                    # Remove trailing dots/leader characters (e.g., "Chapter 1. . . . ." -> "Chapter 1")
                    # Pattern matches dots with spaces between them: ". . . . ." or "....."
                    text_clean = re.sub(r'[\.\s]+$', '', text_clean).strip()
                    
                    if text_clean and text_clean not in toc:  # Avoid duplicates
                        # Skip standalone numbers (likely page numbers or section markers)
                        if re.match(r'^\d+$', text_clean):
                            continue
                        level = _infer_toc_level_from_text(text_clean)
                        toc[text_clean] = level
                        list_count += 1
        
        logger.debug(f"Extracted {list_count} entries from List Paragraph style")
        logger.info(f"Total unique TOC entries extracted: {len(toc)}")

        logger.debug(f"Extracted tocs : {toc}")
        
        return toc
        
    except Exception as e:
        logger.error(f"Error extracting combined TOC from {docx_path}: {e}")
        return {}


def _infer_toc_level_from_text(text: str) -> int:
    """
    Infer TOC level from text content.
    
    Heuristics:
    - "Chapter X" -> level 1
    - "X " (single number) -> level 2
    - "X.Y " (two numbers) -> level 3
    - "X.Y.Z " (three numbers) -> level 4
    - Common sections -> level 1
    - Default -> level 2
    
    Args:
        text: TOC entry text
        
    Returns:
        Inferred level (1-5)
    """
    # Check for chapter
    if text.lower().startswith('chapter '):
        return 1
    
    # Check for numbered sections (e.g., "1.1", "1.2.3")
    match = re.match(r'^(\d+(?:\.\d+)*)\s+', text)
    if match:
        dots = match.group(1).count('.')
        # Correct mapping: "1" (0 dots) -> 2, "1.1" (1 dot) -> 2, "1.2.1" (2 dots) -> 3
        # This ensures "1.1" and "1.2" are at the same level as their parent "1"
        return min(max(dots, 1) + 1, 5)
    
    # Check for common top-level sections
    if text in ['Preface', 'Introduction', 'Contents', 'Notices', 'Trademarks', 'Appendix']:
        return 1
    
    # Default to level 2 for other entries
    return 2
